"""
md_a_word.py — Convierte evaluaciones Markdown a Word (.docx)
Uso: python md_a_word.py "ruta/al/archivo.md"
     python md_a_word.py "ruta/al/archivo.md" --sin-ia
"""

import sys, re, argparse, subprocess
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────
# CONFIGURACION
# ─────────────────────────────────────────────
CLAUDE_EXE     = r"C:\Users\braya\AppData\Roaming\Claude\claude-code\2.1.138\claude.exe"
ESCUELA        = 'ESCUELA DE EDUCACION BASICA "CARIBE"'
ESCUELA_RENDER = 'ESCUELA DE EDUCACIÓN BÁSICA "CARIBE"'
ACUERDO        = "ACUERDO MINISTERIAL # 020-12 DE JULIO DEL 2013"
UBICACION      = "GARCIA MORENO   -   ORELLANA   -   ECUADOR"
LOGO_IZQUIERDO = None   # r"C:\...\logo_caribe.png"
COLOR_AZUL     = RGBColor(0x1F, 0x35, 0x64)   # azul oscuro institucional
COLOR_SECCION  = RGBColor(0x1F, 0x6F, 0xBF)   # azul medio para secciones
COLOR_PUNTOS   = RGBColor(0x1F, 0x6F, 0xBF)   # azul para "(X punto)"

# ─────────────────────────────────────────────
# LATEX → TEXTO LEGIBLE
# ─────────────────────────────────────────────
LATEX_MAP = {
    r'\cdot': '×', r'\times': '×', r'\div': '÷',
    r'\rightarrow': '→', r'\leftarrow': '←',
    r'\leq': '≤', r'\geq': '≥', r'\neq': '≠',
    r'\approx': '≈', r'\sqrt': '√', r'\pm': '±',
    r'\text{': '', r'\frac': '/',
}
SUPERSCRIPTS = str.maketrans('0123456789', '⁰¹²³⁴⁵⁶⁷⁸⁹')

def latex_a_texto(texto):
    texto = re.sub(r'\$\$(.+?)\$\$', lambda m: _conv(m.group(1)), texto, flags=re.DOTALL)
    texto = re.sub(r'\$(.+?)\$',     lambda m: _conv(m.group(1)), texto)
    return texto

def _conv(expr):
    expr = expr.strip()
    for k, v in LATEX_MAP.items():
        expr = expr.replace(k, v)
    expr = re.sub(r'\^\{(\d+)\}', lambda m: m.group(1).translate(SUPERSCRIPTS), expr)
    expr = re.sub(r'\^(\d)',       lambda m: m.group(1).translate(SUPERSCRIPTS), expr)
    return expr.replace('{','').replace('}','').replace('\\','').strip()

# ─────────────────────────────────────────────
# REESCRITURA CON CLAUDE CODE CLI
# ─────────────────────────────────────────────
PROMPT_IA = """Eres un editor de evaluaciones escolares de matematicas para Ecuador (educacion basica).
Reescribe el siguiente texto de evaluacion siguiendo estas reglas ESTRICTAS:

1. Usa SIEMPRE "usted" para dirigirte al estudiante (nunca "tu").
2. Lenguaje simple y claro para estudiantes de educacion basica.
3. ELIMINA COMPLETAMENTE cualquier referencia a: "Profe Alex", "Alex", "el profesor",
   "segun el video", "como se explico", "segun se indico", "explicada por", "indicada por",
   "vista en clase" o cualquier frase que mencione una fuente externa. Reemplaza esa parte
   de la oracion con lenguaje institucional neutro. Por ejemplo:
   - "segun la jerarquia explicada por el Profe Alex" → "segun la jerarquia de las operaciones"
   - "como se vio en el video" → (eliminar completamente esa frase)
4. NO cambies: formulas matematicas, nombres de secciones, puntajes, opciones a/b/c/d,
   tablas, bancos de palabras, lineas de respuesta (____), numeros de preguntas.
5. Mantén el mismo numero de preguntas y secciones.
6. Devuelve SOLO el texto reescrito, sin explicaciones ni comentarios."""

# Filtro de respaldo: elimina referencias al Profe Alex aunque la IA falle
PATRON_PROFE = re.compile(
    r'\s*(explicada?\s+por\s+el\s+Profe\s+Alex|'
    r'seg[uú]n\s+el\s+Profe\s+Alex|'
    r'el\s+Profe\s+Alex|Profe\s+Alex|'
    r'seg[uú]n\s+el\s+video|'
    r'como\s+se\s+explic[oó]|'
    r'como\s+se\s+indic[oó])',
    re.IGNORECASE
)

def reescribir_con_ia(texto_md):
    print("  Reescribiendo con Claude Code...")
    try:
        prompt = f"{PROMPT_IA}\n\nEVALUACION:\n\n{texto_md}"
        r = subprocess.run(
            [CLAUDE_EXE, "--print"],
            input=prompt, capture_output=True,
            text=True, encoding="utf-8", timeout=120
        )
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout.strip()
        print(f"  Advertencia IA: {r.stderr[:150]}\n  Continuando sin reescritura.")
    except Exception as e:
        print(f"  Error IA: {e}\n  Continuando sin reescritura.")
    return texto_md

# ─────────────────────────────────────────────
# CALCULAR PUNTOS POR PREGUNTA
# ─────────────────────────────────────────────
def calcular_puntos(texto):
    """Devuelve dict {num_pregunta_original: pts_float}.
    Detecta preguntas con Q1/Q2 o 1./1) dentro de cada sección."""
    pts = {}
    for seccion in re.split(r'###\s+SECCI[OÓ]N', texto)[1:]:
        m = re.search(r'\(?\$?(\d+(?:\.\d+)?)\$?\s*puntos?\)?', seccion, re.IGNORECASE)
        if not m:
            continue
        total = float(m.group(1))
        # Buscar Q1, Q2 ...
        preguntas = re.findall(r'\bQ(\d+)[.).\*\s]', seccion)
        if not preguntas:
            # Buscar 1. o 1) al inicio de línea
            preguntas = re.findall(r'(?m)^\*?\*?(\d+)[.)]\s', seccion)
        if preguntas:
            c = round(total / len(preguntas), 2)
            for q in preguntas:
                pts[int(q)] = c
    return pts

# ─────────────────────────────────────────────
# HELPERS WORD
# ─────────────────────────────────────────────
def set_run(run, nombre='Arial', tam=11, bold=False, color=None, italic=False):
    run.font.name = nombre
    run.font.size = Pt(tam)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def new_para(doc, alin=WD_ALIGN_PARAGRAPH.LEFT, antes=0, despues=3):
    p = doc.add_paragraph()
    p.alignment = alin
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after  = Pt(despues)
    return p

def agregar_texto_mixto(p, texto, tam=10.5, bold_base=False, color=None):
    """Parsea **negrita** inline y agrega runs al párrafo."""
    partes = re.split(r'(\*\*[^*]+\*\*)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            r = p.add_run(parte[2:-2])
            set_run(r, tam=tam, bold=True, color=color)
        else:
            r = p.add_run(parte)
            set_run(r, tam=tam, bold=bold_base, color=color)

def limpiar_negrita(texto):
    """Quita marcadores ** del texto para celdas de tabla."""
    return re.sub(r'\*\*([^*]+)\*\*', r'\1', texto)

# ─────────────────────────────────────────────
# CUADRÍCULA TIPO PAPEL CUADRICULADO
# ─────────────────────────────────────────────
def agregar_cuadricula(doc, filas=4, columnas=30, tam_celda_cm=0.5):
    """Inserta una tabla con celdas pequeñas tipo papel cuadriculado."""
    tabla = doc.add_table(rows=filas, cols=columnas)
    tabla.autofit = False

    ancho    = int(Cm(tam_celda_cm))
    # Altura en twips: 1 cm = 567 twips
    alto_twips = int(tam_celda_cm * 567)

    for fi in range(filas):
        # Altura de fila en twips
        tr = tabla.rows[fi]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(alto_twips))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

        for ci in range(columnas):
            cel = tabla.cell(fi, ci)
            cel.width = ancho

            # Bordes finos gris claro
            tc   = cel._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for lado in ('top', 'left', 'bottom', 'right'):
                borde = OxmlElement(f'w:{lado}')
                borde.set(qn('w:val'),   'single')
                borde.set(qn('w:sz'),    '4')       # grosor fino
                borde.set(qn('w:space'), '0')
                borde.set(qn('w:color'), 'AAAAAA')  # gris claro
                tcBorders.append(borde)
            tcPr.append(tcBorders)

            # Celda vacía, fuente mínima
            p = cel.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            r = p.add_run(' ')
            r.font.size = Pt(4)

    new_para(doc, despues=3)

# ─────────────────────────────────────────────
# ENCABEZADO INSTITUCIONAL
# ─────────────────────────────────────────────
def agregar_encabezado(doc):
    cols = 2 if LOGO_IZQUIERDO else 1
    tabla = doc.add_table(rows=1, cols=cols)
    tabla.autofit = False

    if LOGO_IZQUIERDO:
        tabla.columns[0].width = Cm(2.5)
        tabla.columns[1].width = Cm(16.0)
        cel_logo = tabla.cell(0, 0)
        cel_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cel_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO_IZQUIERDO, width=Cm(2.2))
        cel_texto = tabla.cell(0, 1)
    else:
        tabla.columns[0].width = Cm(18.5)
        cel_texto = tabla.cell(0, 0)

    cel_texto.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    datos = [
        (ESCUELA_RENDER, 13, True),
        (ACUERDO,         9, False),
        (UBICACION,       9, False),
    ]
    # Quitar párrafo vacío inicial
    for i, (txt, tam, neg) in enumerate(datos):
        if i == 0:
            p = cel_texto.paragraphs[0]
        else:
            p = cel_texto.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(txt)
        set_run(r, tam=tam, bold=neg, color=COLOR_AZUL)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────
# DATOS DEL ESTUDIANTE
# ─────────────────────────────────────────────
def agregar_datos_estudiante(doc, titulo):
    # Título evaluación
    p = new_para(doc, WD_ALIGN_PARAGRAPH.LEFT, antes=4, despues=6)
    r = p.add_run(titulo)
    set_run(r, tam=12, bold=True, color=COLOR_AZUL)

    # Nombre
    p = new_para(doc, despues=2)
    r = p.add_run("Nombre: ")
    set_run(r, tam=11, bold=True)
    r2 = p.add_run("_" * 50)
    set_run(r2, tam=11)

    # Curso + Fecha en tabla sin bordes
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Cm(10)
    t.columns[1].width = Cm(8.5)
    for cel, txt in [(t.cell(0,0), "Curso: ________________________"),
                     (t.cell(0,1), "Fecha: ________________")]:
        p2 = cel.paragraphs[0]
        r3 = p2.add_run(txt)
        set_run(r3, tam=11, bold=True if "Curso" in txt or "Fecha" in txt else False)
        # Bold solo la etiqueta
        p2.clear()
        for parte, neg in [("Curso: " if "Curso" in txt else "Fecha: ", True),
                            ("________________________" if "Curso" in txt else "________________", False)]:
            rx = p2.add_run(parte)
            set_run(rx, tam=11, bold=neg)

    doc.add_paragraph().paragraph_format.space_after = Pt(1)

    # Puntaje
    p = new_para(doc, despues=8)
    r = p.add_run("Puntaje: ")
    set_run(r, tam=11, bold=True)
    r2 = p.add_run("_____ / 10 puntos")
    set_run(r2, tam=11)

# ─────────────────────────────────────────────
# PARSEO MARKDOWN → BLOQUES
# ─────────────────────────────────────────────
def parsear(texto):
    bloques = []
    lineas = texto.split('\n')
    i = 0
    while i < len(lineas):
        L = lineas[i]

        if L.startswith('# '):
            bloques.append({'t': 'titulo', 'v': L[2:].strip()})

        elif L.startswith('### '):
            v = L[4:].strip()
            # Quitar prefijo "SECCIÓN X: "
            v = re.sub(r'^SECCI[OÓ]N\s+[IVXLC\d]+[:\.\-\s]+', '', v, flags=re.IGNORECASE).strip()
            # Quitar puntajes al final: "(2 puntos)", "($3$ puntos)", etc.
            v = re.sub(r'\s*\(?\$?\d+\$?\s*puntos?\)?\.?\s*$', '', v, flags=re.IGNORECASE).strip()
            bloques.append({'t': 'seccion', 'v': v})

        elif L.startswith('|'):
            filas = []
            while i < len(lineas) and lineas[i].startswith('|'):
                if not re.match(r'^\|[\s\-\|:]+\|$', lineas[i]):
                    celdas = [limpiar_negrita(c.strip()) for c in lineas[i].strip('|').split('|')]
                    filas.append(celdas)
                i += 1
            bloques.append({'t': 'tabla', 'filas': filas})
            continue

        elif re.match(r'^-{3,}$', L.strip()):
            pass  # ignorar separadores

        elif L.strip():
            ls = L.strip()
            # ── Ignorar bloque "Identificación:" y sus viñetas ──────────
            # Cabecera: **Identificación:** o Identificación:
            if re.match(r'^\*?\*?Identificaci[oó]n\*?\*?:', ls, re.IGNORECASE):
                pass
            # Viñetas con datos del estudiante: *  **Nombre:** / - Nombre: / * Nombre:
            elif re.match(r'^[\*\-]\s+\*?\*?(Nombre|Curso|Fecha|Puntaje)\*?\*?:', ls, re.IGNORECASE):
                pass
            # Líneas sueltas de datos del estudiante sin viñeta
            elif re.match(r'^\*?\*?(Nombre|Curso|Fecha|Puntaje)\*?\*?[:\*]', ls, re.IGNORECASE):
                pass
            # Instrucciones generales de identificación (línea que dice "complete sus datos...")
            elif re.match(r'^\*?\*?Instrucciones?\s+Generales?\*?\*?:', ls, re.IGNORECASE):
                pass
            else:
                bloques.append({'t': 'parrafo', 'v': ls})

        i += 1
    return bloques

# ─────────────────────────────────────────────
# NUMERACION DE PREGUNTAS
# ─────────────────────────────────────────────
def detectar_pregunta(texto):
    """Retorna (num_original, texto_limpio) si la línea es una pregunta, sino (None, texto).
    Detecta todos los patrones posibles:
      Q1. | Q1) | **Q1.** | 1. | 1) | **1.** | 1.  **Q1 (Tema):** texto"""
    # Patrón mixto: número + Q juntos  →  "1.  **Q1 (Tema):** texto"
    m = re.match(r'^\*?\*?(\d+)[.)]\s*\*?\*?Q\d+[^:]*:\*?\*?\s*', texto)
    if m:
        n = int(m.group(1))
        resto = texto[m.end():].strip().lstrip('*').strip()
        return n, resto
    # Patrón solo Q: Q1. | Q1) | **Q1.**
    m = re.match(r'^\*?\*?Q(\d+)[\.)\*\s]', texto)
    if m:
        n = int(m.group(1))
        resto = texto[m.end():].strip().lstrip('*').strip()
        return n, resto
    # Patrón numérico: 1. | 1) | **1.**  (seguido de texto, no de opción a/b/c/d)
    m = re.match(r'^\*?\*?(\d+)[.)]\*?\*?\s+', texto)
    if m:
        n = int(m.group(1))
        resto = texto[m.end():].strip().lstrip('*').strip()
        # Evitar que opciones tipo "a) texto" sean detectadas como pregunta 97
        if not re.match(r'^[a-dA-D][.)]\s', resto) and n <= 50:
            return n, resto
    return None, texto

# ─────────────────────────────────────────────
# INSTRUCCIONES PREDETERMINADAS POR SECCIÓN
# Se insertan SOLO si la sección no tiene ya una línea "Instrucción:"
# ─────────────────────────────────────────────
INSTRUCCIONES_DEFAULT = {
    # Selección múltiple
    r'selecci[oó]n\s+m[uú]ltiple':
        'Lea atentamente cada pregunta y marque la opción correcta en el espacio indicado.',
    r'opci[oó]n\s+m[uú]ltiple':
        'Lea atentamente cada pregunta y marque la opción correcta en el espacio indicado.',
    # Verdadero / Falso
    r'verdadero\s+[yo]\s+falso':
        'Escriba V si la afirmación es verdadera o F si es falsa en el espacio indicado.',
    r'verdadero\s*/\s*falso':
        'Escriba V si la afirmación es verdadera o F si es falsa en el espacio indicado.',
    # Completar espacios
    r'completar?\s+espacios?':
        'Complete los espacios en blanco con la palabra o valor correcto.',
    r'llenar\s+espacios?':
        'Complete los espacios en blanco con la palabra o valor correcto.',
    # Relacionar columnas
    r'relacionar?\s+columnas?':
        'Una con una línea cada elemento de la columna A con su correspondiente en la columna B.',
    r'emparejar?':
        'Una con una línea cada elemento de la columna A con su correspondiente en la columna B.',
    # Resolver / ejercicio
    r'resolver?\s+ejercicio':
        'Resuelva el siguiente ejercicio mostrando el procedimiento completo.',
    r'desarrollo':
        'Resuelva el siguiente ejercicio mostrando el procedimiento completo.',
    r'resoluci[oó]n':
        'Resuelva el siguiente ejercicio mostrando el procedimiento completo.',
    # Aplicación
    r'aplicaci[oó]n':
        'Lea el problema, plantee la expresión matemática y resuelva mostrando el procedimiento.',
    r'problema\s+aplicado':
        'Lea el problema, plantee la expresión matemática y resuelva mostrando el procedimiento.',
}

def instruccion_para_seccion(nombre_seccion):
    """Devuelve la instrucción predeterminada para el nombre de sección, o None."""
    nombre_lower = nombre_seccion.lower()
    for patron, instruccion in INSTRUCCIONES_DEFAULT.items():
        if re.search(patron, nombre_lower):
            return instruccion
    return None

# ─────────────────────────────────────────────
# GENERAR WORD
# ─────────────────────────────────────────────
def generar_word(texto_md, ruta_salida):
    doc = Document()

    # Margenes
    for sec in doc.sections:
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    pts_map = calcular_puntos(texto_md)
    texto   = latex_a_texto(texto_md)
    bloques = parsear(texto)

    # Extraer título
    titulo = "EVALUACIÓN"
    for b in bloques:
        if b['t'] == 'titulo':
            titulo = b['v']
            break

    agregar_encabezado(doc)
    agregar_datos_estudiante(doc, titulo)

    num_pregunta = 0  # contador secuencial para renumerar

    for idx, b in enumerate(bloques):
        tipo = b['t']

        if tipo == 'titulo':
            pass  # ya usado

        elif tipo == 'seccion':
            p = new_para(doc, antes=8, despues=3)
            r = p.add_run(b['v'])
            set_run(r, tam=11, bold=True, color=COLOR_SECCION)

            # ¿El siguiente bloque ya trae "Instrucción:"?
            siguiente_es_instruccion = (
                idx + 1 < len(bloques)
                and bloques[idx + 1]['t'] == 'parrafo'
                and re.match(r'^\*?\*?Instrucci[oó]n', bloques[idx + 1]['v'], re.IGNORECASE)
            )
            if not siguiente_es_instruccion:
                instruccion = instruccion_para_seccion(b['v'])
                if instruccion:
                    pi = new_para(doc, despues=3)
                    r_lbl = pi.add_run("Instrucción: ")
                    set_run(r_lbl, tam=10.5, bold=True, italic=True)
                    r_txt = pi.add_run(instruccion)
                    set_run(r_txt, tam=10.5, italic=True)

        elif tipo == 'tabla':
            filas = b['filas']
            if not filas:
                continue
            cols = max(len(f) for f in filas)
            tw = doc.add_table(rows=len(filas), cols=cols)
            tw.style = 'Table Grid'
            for fi, fila in enumerate(filas):
                for ci in range(cols):
                    txt_cel = fila[ci] if ci < len(fila) else ''
                    cel = tw.cell(fi, ci)
                    p2  = cel.paragraphs[0]
                    r2  = p2.add_run(txt_cel)
                    set_run(r2, tam=10, bold=(fi == 0))
            new_para(doc, despues=4)

        elif tipo == 'parrafo':
            v = b['v']

            # ¿Es una pregunta?
            nq, resto = detectar_pregunta(v)
            if nq is not None:
                num_pregunta += 1
                pts = pts_map.get(nq, '')
                # Siempre mostrar puntaje: si no hay dato calculado, asumir 1 punto
                if pts:
                    pts_str = f"({pts} pt)" if pts != 1.0 else "(1 punto)"
                else:
                    pts_str = "(1 punto)"

                p = new_para(doc, antes=4, despues=2)
                # Número con paréntesis: 1) 2) 3)
                r = p.add_run(f"{num_pregunta})  ")
                set_run(r, tam=10.5, bold=True)
                # Texto de la pregunta
                agregar_texto_mixto(p, resto, tam=10.5, bold_base=True)
                # Puntos en color — SIEMPRE
                r2 = p.add_run(f"  {pts_str}")
                set_run(r2, tam=10.5, bold=True, color=COLOR_PUNTOS)

            # ¿Es "Respuesta: ___"?
            elif re.match(r'^\*?\*?Respuesta', v, re.IGNORECASE):
                p = new_para(doc, despues=4)
                r = p.add_run(limpiar_negrita(v))
                set_run(r, tam=10.5, bold=True)

            # ¿Es "Procedimiento"? → reemplazar líneas por cuadrícula
            elif re.match(r'^\*?\*?Procedimiento', v, re.IGNORECASE):
                p = new_para(doc, despues=2)
                r = p.add_run(re.sub(r'_{3,}', '', limpiar_negrita(v)).strip())
                set_run(r, tam=10.5, bold=True)
                agregar_cuadricula(doc, filas=6, columnas=40, tam_celda_cm=0.32)

            # ¿Es "Expresión matemática: ___"?
            elif re.match(r'^\*?\*?Expresi[oó]n', v, re.IGNORECASE):
                p = new_para(doc, despues=2)
                r = p.add_run(re.sub(r'_{3,}', '', limpiar_negrita(v)).strip() + '  ')
                set_run(r, tam=10.5, bold=True)

            # ¿Son líneas de guiones vacías? → ignorar (reemplazadas por cuadrícula)
            elif re.match(r'^_{4,}$', v.strip()):
                pass

            # ¿Es "Instrucción:"?
            elif re.match(r'^\*?\*?Instrucci[oó]n', v, re.IGNORECASE):
                p = new_para(doc, despues=2)
                agregar_texto_mixto(p, limpiar_negrita(v), tam=10.5)

            # ¿Es opción a) b) c) d)?
            elif re.match(r'^[a-d]\)', v):
                p = new_para(doc, despues=1)
                r = p.add_run(limpiar_negrita(v))
                set_run(r, tam=10.5)
                p.paragraph_format.left_indent = Cm(0.8)

            # Cualquier otra línea
            else:
                p = new_para(doc, despues=2)
                agregar_texto_mixto(p, v, tam=10.5)

    doc.save(str(ruta_salida))
    print(f"  Word guardado: {ruta_salida}")

# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────
def procesar_archivo(ruta_md, sin_ia=False):
    """Convierte un archivo .md de evaluacion a .docx. Retorna True si tuvo exito."""
    nombre_limpio = re.sub(r'[^\w\s\-]', '', ruta_md.stem).strip()
    nombre_limpio = re.sub(r'\s+', ' ', nombre_limpio)
    ruta_salida   = ruta_md.parent / (nombre_limpio + '.docx')

    if ruta_salida.exists():
        print(f"  [OMITIDO] Ya existe: {ruta_salida.name}")
        return False

    print(f"\n{'='*50}")
    print(f"  {ruta_md.name}")
    print(f"{'='*50}\n")

    texto = ruta_md.read_text(encoding='utf-8')

    if not sin_ia:
        texto = reescribir_con_ia(texto)
    else:
        print("  Modo rapido: sin reescritura IA")

    # Filtro de respaldo: eliminar referencias al Profe Alex
    texto = PATRON_PROFE.sub('', texto)

    print("  Generando Word...")
    generar_word(texto, ruta_salida)
    print(f"  Listo: {ruta_salida.name}\n")
    return True


def main():
    parser = argparse.ArgumentParser(
        description='Convierte evaluaciones Markdown a Word (.docx)'
    )
    parser.add_argument('archivo', nargs='?', help='Ruta al archivo .md (o carpeta base con --todos)')
    parser.add_argument('--sin-ia', action='store_true', help='Sin reescritura IA (modo rapido)')
    parser.add_argument('--todos', metavar='CARPETA',
                        help='Procesar todos los archivos Evaluaci*.md dentro de CARPETA')
    args = parser.parse_args()

    # ── MODO --todos ──────────────────────────────────────────────────
    if args.todos:
        base = Path(args.todos)
        if not base.exists():
            print(f"Carpeta no encontrada: {base}"); sys.exit(1)

        # Buscar archivos que contengan "Evaluaci" en el nombre (con o sin emoji)
        patrones = ['**/*Evaluaci*.md', '**/*evaluaci*.md']
        encontrados = []
        for pat in patrones:
            encontrados.extend(base.glob(pat))
        # Eliminar duplicados y ordenar
        encontrados = sorted(set(encontrados))

        if not encontrados:
            print(f"No se encontraron evaluaciones en: {base}"); sys.exit(0)

        print(f"\nEncontradas {len(encontrados)} evaluaciones en: {base}\n")
        convertidos, omitidos = 0, 0
        for ruta in encontrados:
            ok = procesar_archivo(ruta, sin_ia=args.sin_ia)
            if ok: convertidos += 1
            else:  omitidos += 1

        print(f"\n{'='*50}")
        print(f"  RESUMEN: {convertidos} convertidos | {omitidos} omitidos")
        print(f"{'='*50}\n")
        return

    # ── MODO archivo individual ───────────────────────────────────────
    if not args.archivo:
        parser.print_help(); sys.exit(1)

    ruta_md = Path(args.archivo)
    if not ruta_md.exists():
        print(f"No se encontro: {ruta_md}"); sys.exit(1)

    procesar_archivo(ruta_md, sin_ia=args.sin_ia)


if __name__ == '__main__':
    main()
